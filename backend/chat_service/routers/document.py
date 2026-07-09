"""文档生成路由 - 从 AI 回复生成 Word 文档"""
import uuid
from pathlib import Path
from datetime import datetime

from fastapi import APIRouter, Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from common.dependencies import get_current_user
from common.config import settings
from common.response import success_response, ErrorCode
from common.exceptions import BusinessError
from common.logging_config import get_logger

logger = get_logger()

router = APIRouter(prefix="/api/chat", tags=["文档生成"])


class GenerateDocRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200, description="文档标题")
    content: str = Field(..., min_length=1, max_length=50000, description="文档正文内容")


def _generate_docx(title: str, content: str, output_path: str):
    """使用 python-docx 生成 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 生成时间
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = None  # 默认颜色

    doc.add_paragraph()  # 空行

    # 正文内容：按段落分割
    paragraphs = content.split("\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # 检测标题行（## 开头或短行）
        if para_text.startswith("## "):
            doc.add_heading(para_text[3:], level=2)
        elif para_text.startswith("# "):
            doc.add_heading(para_text[2:], level=1)
        elif para_text.startswith("- ") or para_text.startswith("* "):
            p = doc.add_paragraph(para_text[2:], style="List Bullet")
        elif para_text.startswith("**") and para_text.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(para_text[2:-2])
            run.bold = True
        else:
            p = doc.add_paragraph(para_text)
            # 设置正文字体
            for run in p.runs:
                run.font.size = Pt(11)

    doc.save(output_path)


@router.post("/sessions/{session_id}/generate-doc")
async def generate_doc(
    session_id: int,
    body: GenerateDocRequest,
    user: dict = Depends(get_current_user),
):
    """从 AI 回复生成 Word 文档"""
    doc_dir = Path(settings.UPLOAD_DIR) / "chat_docs"
    doc_dir.mkdir(parents=True, exist_ok=True)

    # 只保留 ASCII 安全字符作为文件名（避免 URL 编码问题）
    safe_title = "".join(c for c in body.title if c.isascii() and (c.isalnum() or c == "_"))
    safe_title = safe_title.strip("_")[:50]
    if not safe_title:
        safe_title = "document"
    file_name = f"{safe_title}_{uuid.uuid4().hex[:8]}.docx"
    output_path = doc_dir / file_name

    try:
        _generate_docx(body.title, body.content, str(output_path))
    except Exception as e:
        logger.error("生成文档失败: %s", e)
        raise BusinessError(ErrorCode.SERVER_ERROR, "文档生成失败，请稍后重试")

    logger.info("文档已生成: session_id=%s, user_id=%s, file=%s",
                session_id, user["user_id"], file_name)

    return success_response({
        "file_name": file_name,
        "title": body.title,
        "download_url": f"/api/chat/documents/{file_name}",
    })


@router.get("/documents/{file_name}")
async def download_doc(
    file_name: str,
    user: dict = Depends(get_current_user),
):
    """下载生成的文档"""
    file_path = Path(settings.UPLOAD_DIR) / "chat_docs" / file_name
    if not file_path.exists():
        raise BusinessError(ErrorCode.NOT_FOUND, "文件不存在或已被清理")

    return FileResponse(
        path=str(file_path),
        filename=file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
